from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
import PyPDF2
from reportlab.pdfgen import canvas
import io
import os
from docx import Document

app = Flask(__name__)

# ----- FIXED UPLOAD PATH HANDLING -----
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER


def process_pdf(filepath, data):

    reader = PyPDF2.PdfReader(filepath)
    writer = PyPDF2.PdfWriter()

    color = data["color"]

    for i, page in enumerate(reader.pages):

        packet = io.BytesIO()

        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)

        c = canvas.Canvas(packet, pagesize=(page_width, page_height))

        # ----- HEADER ONLY ON FIRST PAGE -----
        if i == 0:
            c.drawString(50, page_height - 30, data["header_left"])

            c.setFont("Helvetica-Bold", 16)
            c.setFillColorRGB(color[0], color[1], color[2])

            c.drawCentredString(
                page_width / 2,
                page_height - 100,
                data["center"]
            )

            c.setFont("Helvetica", 12)
            c.setFillColorRGB(0, 0, 0)

            c.drawRightString(page_width - 50, page_height - 30, data["header_right"])

        # ----- FOOTER ON EVERY PAGE -----
        c.drawString(50, 25, data["footer_left"])
        c.drawRightString(page_width - 50, 25, data["footer_right"])

        if data["pageno"] == "yes":
            c.drawCentredString(page_width / 2, 25, f"Page {i+1}")

        c.save()

        packet.seek(0)
        overlay = PyPDF2.PdfReader(packet)

        # ----- SHIFT ONLY FIRST PAGE (NO CROPPING) -----
        if i == 0:
            page.add_transformation(
                PyPDF2.Transformation().translate(0, -60)
            )

        page.merge_page(overlay.pages[0])
        writer.add_page(page)

    name_only = os.path.splitext(os.path.basename(filepath))[0]
    output_path = os.path.join(UPLOAD_FOLDER, name_only + "_edited.pdf")

    with open(output_path, "wb") as output:
        writer.write(output)

    return output_path


@app.route("/", methods=["GET", "POST"])
def index():

    if request.method == "POST":

        file = request.files["file"]

        # ----- SAFE FILENAME HANDLING -----
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)

        file.save(filepath)

        r = int(request.form.get("r", 255)) / 255
        g = int(request.form.get("g", 0)) / 255
        b = int(request.form.get("b", 0)) / 255

        data = {
            "header_left": request.form["header_left"],
            "header_right": request.form["header_right"],
            "center": request.form["center"],
            "footer_left": request.form["footer_left"],
            "footer_right": request.form["footer_right"],
            "pageno": request.form.get("pageno", "no"),
            "color": (r, g, b)
        }

        ext = filename.split(".")[-1].lower()

        if ext == "pdf":
            output = process_pdf(filepath, data)
            return send_file(output, as_attachment=True)

        elif ext == "docx":

            doc = Document(filepath)

            section = doc.sections[0]
            section.header.paragraphs[0].text = f'{data["header_left"]} {data["center"]} {data["header_right"]}'
            section.footer.paragraphs[0].text = f'{data["footer_left"]} {data["footer_right"]}'

            name_only = os.path.splitext(filename)[0]
            output = os.path.join(UPLOAD_FOLDER, name_only + "_edited.docx")

            doc.save(output)
            return send_file(output, as_attachment=True)

    return render_template("index.html")


if __name__ == "__main__":
    app.run(debug=True)
